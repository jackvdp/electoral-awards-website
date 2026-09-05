#!/usr/bin/env python3
"""Read and edit .docx documents while preserving their formatting.

Subcommands:
  text      Dump every paragraph with its index, so edits can be targeted
  blocks    List body blocks (paragraphs and tables) in document order
  replace   Find-and-replace across body, tables, headers and footers
  set       Overwrite the whole text of one paragraph, keeping its style
  delete    Delete a range of paragraphs
  insert    Insert new paragraphs after a given index, cloning another paragraph's style
  truncate  Cut everything after a block, tables included, keeping the letterhead

Every editing subcommand prints a report of what actually changed, and refuses to write
a file where nothing matched, so a silent no-op cannot be mistaken for success.
"""

import argparse
import copy
import json
import sys

import docx


# --------------------------------------------------------------------------------------
# Paragraph access
# --------------------------------------------------------------------------------------

def body_paragraphs(doc):
    """Body paragraphs only. These are the ones `set`, `delete` and `insert` address."""
    return doc.paragraphs


def all_paragraphs(doc):
    """Every paragraph in the document, including tables, headers and footers.

    `replace` walks this so that dates in a running header get corrected too. Text boxes
    are not reachable through python-docx's object model and are handled separately.
    """
    for p in doc.paragraphs:
        yield ("body", p)
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    yield ("table", p)
    for s in doc.sections:
        parts = (
            ("header", s.header),
            ("footer", s.footer),
            ("first-page header", s.first_page_header),
            ("first-page footer", s.first_page_footer),
            ("even-page header", s.even_page_header),
            ("even-page footer", s.even_page_footer),
        )
        for label, part in parts:
            if part is None:
                continue
            for p in part.paragraphs:
                yield (label, p)


# --------------------------------------------------------------------------------------
# Run-aware replacement
# --------------------------------------------------------------------------------------

def replace_in_paragraph(paragraph, old, new):
    """Replace `old` with `new`, coping with text split across runs.

    Word fragments a paragraph into runs at every formatting change, and often at
    arbitrary points besides ("12th February 2026" is stored as "12th " + "February 2026").
    A naive per-run replace therefore misses any match spanning a boundary, which is the
    classic way this kind of script silently does nothing.

    Matches inside a single run are edited in place, so inline bold or italic elsewhere in
    the paragraph survives. Only when a match spans runs are those runs collapsed, and
    then just the ones the match actually touches.
    """
    runs = paragraph.runs
    if not runs:
        return 0

    full = "".join(r.text for r in runs)
    if old not in full:
        return 0

    # Character span covered by each run.
    spans, pos = [], 0
    for r in runs:
        spans.append((pos, pos + len(r.text)))
        pos += len(r.text)

    # Collect match positions, then apply right to left so earlier offsets stay valid.
    starts, at = [], full.find(old)
    while at != -1:
        starts.append(at)
        at = full.find(old, at + len(old))

    for start in reversed(starts):
        end = start + len(old)
        touched = [i for i, (s, e) in enumerate(spans) if s < end and e > start]
        first = touched[0]
        fs, _ = spans[first]

        if len(touched) == 1:
            r = runs[first]
            off = start - fs
            r.text = r.text[:off] + new + r.text[off + len(old):]
        else:
            # Head of the first run, the replacement, then the tail of the last run.
            last = touched[-1]
            ls, _ = spans[last]
            head = runs[first].text[:start - fs]
            tail = runs[last].text[end - ls:]
            runs[first].text = head + new + tail
            for i in touched[1:]:
                runs[i].text = ""

        # Offsets shifted, so recompute before the next match.
        spans, pos = [], 0
        for r in runs:
            spans.append((pos, pos + len(r.text)))
            pos += len(r.text)

    return len(starts)


def set_paragraph_text(paragraph, text):
    """Replace a paragraph's entire text, keeping the first run's character formatting."""
    if not paragraph.runs:
        paragraph.add_run(text)
        return
    paragraph.runs[0].text = text
    for r in paragraph.runs[1:]:
        r.text = ""


# --------------------------------------------------------------------------------------
# Subcommands
# --------------------------------------------------------------------------------------

def cmd_text(args):
    doc = docx.Document(args.file)
    for i, p in enumerate(body_paragraphs(doc)):
        if not args.all and not p.text.strip():
            continue
        style = p.style.name if p.style is not None else "?"
        line = f"[{i:>3}] ({style}) {p.text}"
        if args.runs:
            line += "\n      runs: " + repr([r.text for r in p.runs])
        print(line)

    extras = [(w, p) for w, p in all_paragraphs(doc) if w != "body" and p.text.strip()]
    if extras:
        print("\n--- outside the body (editable via `replace`, not by index) ---")
        for where, p in extras:
            print(f"[{where}] {p.text}")


def iter_blocks(doc):
    """Yield ('paragraph'|'table', object) for each body-level block, in document order.

    python-docx exposes doc.paragraphs and doc.tables as separate flat lists, which loses
    their interleaving. Anything that needs to know what follows what (truncating an
    appendix, keeping a signature table) has to walk the XML instead.
    """
    from docx.table import Table
    from docx.text.paragraph import Paragraph

    body = doc.element.body
    for child in body.iterchildren():
        if child.tag.endswith('}p'):
            yield ("paragraph", Paragraph(child, doc))
        elif child.tag.endswith('}tbl'):
            yield ("table", Table(child, doc))


def cmd_blocks(args):
    doc = docx.Document(args.file)
    for i, (kind, obj) in enumerate(iter_blocks(doc)):
        if kind == "paragraph":
            text = obj.text.strip()
            if not text and not args.all:
                continue
            print(f"[{i:>3}] paragraph  {text[:90]}")
        else:
            first = ""
            if obj.rows and obj.rows[0].cells:
                first = obj.rows[0].cells[0].text.strip()[:40]
            print(f"[{i:>3}] TABLE      {len(obj.rows)}x{len(obj.columns)}  first cell: {first!r}")


def cmd_truncate(args):
    doc = docx.Document(args.file)
    blocks = list(iter_blocks(doc))
    if args.after >= len(blocks):
        sys.exit(f"Error: block {args.after} does not exist (document has {len(blocks)})")

    body = doc.element.body
    keep_el = blocks[args.after][1]._element
    removed = 0
    seen_keep = False
    for child in list(body.iterchildren()):
        if child is keep_el:
            seen_keep = True
            continue
        if not seen_keep:
            continue
        # sectPr carries page setup and the header reference. Removing it would strip the
        # letterhead from the document, so it always stays.
        if child.tag.endswith('}sectPr'):
            continue
        body.remove(child)
        removed += 1

    out = args.out or args.file
    doc.save(out)
    print(f"removed {removed} block(s) after [{args.after}], written to {out}")


def cmd_replace(args):
    if args.edits:
        with open(args.edits) as fh:
            edits = json.load(fh)
    else:
        if args.old is None or args.new is None:
            sys.exit("Error: give --edits FILE.json, or both --old and --new")
        edits = [{"old": args.old, "new": args.new}]

    doc = docx.Document(args.file)
    total, report = 0, []
    for e in edits:
        old, new = e["old"], e["new"]
        n = sum(replace_in_paragraph(p, old, new) for _, p in all_paragraphs(doc))
        report.append((n, old, new))
        total += n

    for n, old, new in report:
        mark = "  ok" if n else "MISS"
        print(f"{mark}  {n} x  {old[:60]!r} -> {new[:60]!r}")

    misses = [r for r in report if r[0] == 0]
    if total == 0:
        sys.exit("\nError: nothing matched, file not written. Run `text` to check the wording.")

    out = args.out or args.file
    doc.save(out)
    print(f"\n{total} replacement(s) written to {out}")
    if misses:
        print(f"WARNING: {len(misses)} edit(s) matched nothing, listed as MISS above.")


def cmd_set(args):
    doc = docx.Document(args.file)
    paras = body_paragraphs(doc)
    if args.index >= len(paras):
        sys.exit(f"Error: paragraph {args.index} does not exist (document has {len(paras)})")
    before = paras[args.index].text
    set_paragraph_text(paras[args.index], args.text)
    out = args.out or args.file
    doc.save(out)
    print(f"[{args.index}] {before!r}\n   -> {args.text!r}\nwritten to {out}")


def cmd_delete(args):
    doc = docx.Document(args.file)
    paras = body_paragraphs(doc)
    start, end = args.start, args.end if args.end is not None else args.start
    if end >= len(paras):
        sys.exit(f"Error: paragraph {end} does not exist (document has {len(paras)})")
    removed = []
    for i in range(end, start - 1, -1):
        removed.append((i, paras[i].text))
        paras[i]._element.getparent().remove(paras[i]._element)
    out = args.out or args.file
    doc.save(out)
    for i, t in reversed(removed):
        print(f"deleted [{i}] {t[:80]!r}")
    print(f"\n{len(removed)} paragraph(s) removed, written to {out}")


def cmd_insert(args):
    doc = docx.Document(args.file)
    paras = body_paragraphs(doc)
    if args.after >= len(paras):
        sys.exit(f"Error: paragraph {args.after} does not exist (document has {len(paras)})")

    # Clone a paragraph so the new text inherits real styling (font, size, justification)
    # rather than whatever the document's default happens to be.
    template_idx = args.style_from if args.style_from is not None else args.after
    template = paras[template_idx]

    anchor = paras[args.after]._element
    for text in reversed(args.text):
        new_el = copy.deepcopy(template._element)
        anchor.addnext(new_el)
        new_para = docx.text.paragraph.Paragraph(new_el, paras[args.after]._parent)
        set_paragraph_text(new_para, text)

    out = args.out or args.file
    doc.save(out)
    for t in args.text:
        print(f"inserted after [{args.after}]: {t[:80]!r}")
    print(f"\n{len(args.text)} paragraph(s) inserted (style from [{template_idx}]), written to {out}")


# --------------------------------------------------------------------------------------

def main():
    ap = argparse.ArgumentParser(description=__doc__, formatter_class=argparse.RawDescriptionHelpFormatter)
    sub = ap.add_subparsers(dest="cmd", required=True)

    # `text` is the usual starting point; `blocks` is for structural edits.

    p = sub.add_parser("text", help="dump paragraphs with indices")
    p.add_argument("file")
    p.add_argument("--runs", action="store_true", help="also show the run split")
    p.add_argument("--all", action="store_true", help="include empty paragraphs")
    p.set_defaults(func=cmd_text)

    p = sub.add_parser("blocks", help="list body blocks (paragraphs and tables) in order")
    p.add_argument("file")
    p.add_argument("--all", action="store_true", help="include empty paragraphs")
    p.set_defaults(func=cmd_blocks)

    p = sub.add_parser("truncate", help="delete every block after this one, tables included")
    p.add_argument("file")
    p.add_argument("--after", type=int, required=True, help="block index from `blocks`")
    p.add_argument("--out")
    p.set_defaults(func=cmd_truncate)

    p = sub.add_parser("replace", help="find and replace throughout the document")
    p.add_argument("file")
    p.add_argument("--edits", help="JSON file: [{\"old\": ..., \"new\": ...}, ...]")
    p.add_argument("--old")
    p.add_argument("--new")
    p.add_argument("--out", help="write here instead of in place")
    p.set_defaults(func=cmd_replace)

    p = sub.add_parser("set", help="overwrite one paragraph's text")
    p.add_argument("file")
    p.add_argument("--index", type=int, required=True)
    p.add_argument("--text", required=True)
    p.add_argument("--out")
    p.set_defaults(func=cmd_set)

    p = sub.add_parser("delete", help="delete a range of paragraphs")
    p.add_argument("file")
    p.add_argument("--start", type=int, required=True)
    p.add_argument("--end", type=int, help="defaults to --start")
    p.add_argument("--out")
    p.set_defaults(func=cmd_delete)

    p = sub.add_parser("insert", help="insert paragraphs after an index")
    p.add_argument("file")
    p.add_argument("--after", type=int, required=True)
    p.add_argument("--text", action="append", required=True, help="repeatable, one per paragraph")
    p.add_argument("--style-from", type=int, help="clone this paragraph's style instead")
    p.add_argument("--out")
    p.set_defaults(func=cmd_insert)

    args = ap.parse_args()
    args.func(args)


if __name__ == "__main__":
    main()
